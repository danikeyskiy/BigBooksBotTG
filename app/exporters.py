import os
from openpyxl import Workbook
from docx import Document

EXPORT_PATH = "tmp"
        
def export_to_xlsx(models: dict) -> str:
    wb = Workbook()

    for idx, (model_name, model) in enumerate(models.items()):
        sheet = wb.active if idx == 0 else wb.create_sheet(title=model_name)
        sheet.title = model_name

        sample = model.select().first()
        if sample:
            sheet.append(list(sample.__data__.keys()))

            for item in model.select():
                sheet.append(list(item.__data__.values()))

    file_path = os.path.join(EXPORT_PATH, "models.xlsx")
    wb.save(file_path)
    
    return file_path

def export_to_docx(models: dict) -> str:
    doc = Document()
    doc.add_heading('Выгруженные данные из БД', 0)

    for model_name, model in models.items():
        doc.add_heading(model_name, level=1)

        sample = model.select().first()
        if sample:
            doc.add_paragraph(' | '.join(sample.__data__.keys()))

            for item in model.select():
                doc.add_paragraph(' | '.join(str(value) for value in item.__data__.values()))

    file_path = os.path.join(EXPORT_PATH, "models.docx")
    doc.save(file_path)

    return file_path